import json
import re
from pathlib import Path
from typing import Any

from fpdf import FPDF

from gemini_chat import ask_gemini
from io import BytesIO
from docx import Document
from docx.shared import Pt


def _load_cv_structure_markdown() -> str:
    template_path = Path(__file__).with_name("cv_structure.md")
    if template_path.exists():
        return template_path.read_text(encoding="utf-8")
    return "Use an ATS-friendly one-page CV structure."


def _extract_json(text: str) -> dict:
    text = text.strip()
    try:
        return json.loads(text)
    except Exception:
        start = text.find("{")
        end = text.rfind("}")
        if start == -1 or end == -1 or end <= start:
            raise ValueError("AI response did not contain valid JSON.")
        snippet = text[start : end + 1]
        return json.loads(snippet)


def _safe_text(value: str) -> str:
    text = value
    replacements = {
        "•": "-",
        "●": "-",
        "◦": "-",
        "–": "-",
        "—": "-",
        "−": "-",
        "’": "'",
        "‘": "'",
        "“": '"',
        "”": '"',
        "…": "...",
        "\u00a0": " ",
        "\u200b": "",
        "Â": "",
        "â€™": "'",
        "â€˜": "'",
        "â€œ": '"',
        "â€\x9d": '"',
        "â€": '"',
        "â€“": "-",
        "â€”": "-",
        "Â·": "-",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    # Remove malformed markdown emphasis markers that can leak as visible asterisks
    text = re.sub(r"\*{3,}", "", text)
    text = re.sub(r"(?<!\*)\*{1}(?!\*)", "", text)

    text = text.encode("latin-1", errors="ignore").decode("latin-1")
    text = re.sub(r"\s+", " ", text).strip()
    # FPDF can fail when a single token is wider than the printable area.
    # Break very long non-space tokens into smaller chunks.
    parts = text.split()
    normalized: list[str] = []
    for part in parts:
        if len(part) <= 45:
            normalized.append(part)
            continue
        chunks = [part[i : i + 45] for i in range(0, len(part), 45)]
        normalized.append(" ".join(chunks))
    return " ".join(normalized)


def _build_education_from_profile(profile_data: dict[str, Any]) -> list[str]:
    education = profile_data.get("education")
    if not isinstance(education, dict):
        return []

    degree = str(education.get("degree") or "").strip()
    institution = str(education.get("institution") or "").strip()
    location = str(education.get("location") or "").strip()
    graduation_date = str(education.get("graduation_date") or "").strip()

    left = " - ".join(part for part in [degree, institution] if part)
    right = ", ".join(part for part in [location, graduation_date] if part)

    if left and right:
        return [f"{left} | {right}"]
    if left:
        return [left]
    if right:
        return [right]
    return []


def _build_training_certifications_from_profile(profile_data: dict[str, Any]) -> list[str]:
    entries = profile_data.get("training_and_certifications")
    if not isinstance(entries, list):
        return []

    formatted: list[str] = []
    for entry in entries:
        if not isinstance(entry, dict):
            continue
        name = str(entry.get("name") or "").strip()
        provider = str(entry.get("provider") or "").strip()
        duration = str(entry.get("duration") or "").strip()
        line = " - ".join(part for part in [name, provider, duration] if part)
        if line:
            formatted.append(line)
    return formatted


def _normalize_highlights(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    if isinstance(value, str) and value.strip():
        return [line.strip() for line in re.split(r"\n+", value) if line.strip()]
    return []


def _build_professional_experience_from_profile(
    profile_data: dict[str, Any],
    ai_experiences: Any,
) -> list[dict[str, Any]]:
    profile_experiences = profile_data.get("professional_experience")
    if not isinstance(profile_experiences, list):
        return []

    ai_list = ai_experiences if isinstance(ai_experiences, list) else []
    merged: list[dict[str, Any]] = []

    for index, profile_entry in enumerate(profile_experiences):
        if not isinstance(profile_entry, dict):
            continue

        title = str(profile_entry.get("title") or "").strip()
        company = str(profile_entry.get("company") or "").strip()
        duration = str(profile_entry.get("duration") or "").strip()
        profile_description = str(profile_entry.get("description") or "").strip()

        ai_entry = ai_list[index] if index < len(ai_list) and isinstance(ai_list[index], dict) else {}
        highlights = _normalize_highlights(ai_entry.get("highlights"))
        if not highlights and profile_description:
            highlights = [profile_description]

        merged.append(
            {
                "title": title,
                "company": company,
                "duration": duration,
                "highlights": highlights,
            }
        )

    return merged


def _postprocess_sections_from_profile(data: dict[str, Any], profile_data: dict[str, Any]) -> dict[str, Any]:
    for required_key in [
        "header",
        "professional_summary",
        "core_skills",
        "professional_experience",
        "personal_projects",
    ]:
        if required_key not in data:
            raise ValueError(f"AI response missing key: {required_key}")

    data["education"] = _build_education_from_profile(profile_data)
    data["training_certifications"] = _build_training_certifications_from_profile(profile_data)
    data["professional_experience"] = _build_professional_experience_from_profile(
        profile_data,
        data.get("professional_experience"),
    )
    return data


def _canonical_cv_core(cv_structure: str, profile_json: str, job_description: str) -> str:
    return (
        "You are an expert resume writer. Your task is to refine the candidate's CV content using BOTH the profile data and the target job description. "
        "Focus on enhancing the impact, readability, and ATS compatibility of the resume. "
        "Important rules to follow:\n"
        "1. For each professional experience in profile_json, KEEP the job title, company, and duration exactly as they are. "
        "   Only rewrite/rephrase the 'highlights' or 'description' bullets to be more impactful and results-oriented.\n"
        "2. For personal projects, keep the project name intact, but refine the description to emphasize measurable results, relevant technologies, and impact. Include tech_stack and concrete outcomes whenever possible.\n"
        "3. Do NOT invent candidate facts that are not present in profile_json. "
        "4. Do NOT generate education or training_certifications sections; those are provided separately.\n"
        "5. Use concise, action-oriented bullets for professional experience and projects. Prefer strong action verbs and varied language. Avoid repeating the same words excessively.\n"
        "6. Prioritize keywords, responsibilities, and skills from the target job description to make the resume ATS-friendly.\n"
        "7. Ensure all content is free from spelling or grammatical errors.\n"
        "8. Quantify achievements wherever possible to demonstrate clear impact.\n"
        "9. Return ONLY valid JSON following the CV structure defined below; do NOT include markdown fences.\n\n"
        f"CV structure instructions (markdown):\n{cv_structure}\n\n"
        f"Candidate profile data (JSON):\n{profile_json}\n\n"
        f"Target job description:\n{job_description}\n\n"
        "Your output should be a refined JSON CV where only the descriptions/highlights of jobs and projects are rewritten for impact, "
        "while all names, titles, companies, project names, and dates remain unchanged."
    )


def _build_cv_prompt(
    cv_structure: str,
    profile_json: str,
    job_description: str,
    prompt_override: str | None = None,
) -> str:
    if not prompt_override or not prompt_override.strip():
        return _canonical_cv_core(cv_structure, profile_json, job_description)

    override = prompt_override.strip()
    prompt = (
        override.replace("{cv_structure}", cv_structure)
        .replace("{profile_json}", profile_json)
        .replace("{job_description}", job_description)
    )

    return (
        f"{prompt}\n\n"
        f"CV structure instructions (markdown):\n{cv_structure}\n\n"
        f"Candidate profile data (JSON):\n{profile_json}\n\n"
        f"Target job description:\n{job_description}"
    )


def build_resume_sections(
    job_description: str,
    model_name: str,
    profile_data: dict[str, Any],
    prompt_override: str | None = None,
    gemini_api_key: str | None = None,
) -> dict:
    cv_structure = _load_cv_structure_markdown()
    profile_json = json.dumps(profile_data, ensure_ascii=False)

    prompt = _build_cv_prompt(cv_structure, profile_json, job_description, prompt_override)


    raw = ask_gemini(prompt, model_name=model_name, api_key=gemini_api_key)
    data = _extract_json(raw)

    return _postprocess_sections_from_profile(data, profile_data)


def build_resume_bundle(
    job_description: str,
    model_name: str,
    profile_data: dict[str, Any],
    *,
    generate_cv: bool,
    generate_cover_letter: bool,
    generate_email_message: bool,
    full_name: str,
    role_title: str,
    company_name: str,
    prompt_override: str | None = None,
    cover_letter_prompt: str | None = None,
    email_message_prompt: str | None = None,
    gemini_api_key: str | None = None,
) -> dict[str, Any]:
    if not (generate_cv or generate_cover_letter or generate_email_message):
        return {"sections": {}, "cover_letter": None, "email_message": None}

    cv_structure = _load_cv_structure_markdown()
    profile_json = json.dumps(profile_data, ensure_ascii=False)

    output_contract = {
        "sections": "CV sections object following CV schema" if generate_cv else None,
        "cover_letter": "string" if generate_cover_letter else None,
        "email_message": {"subject": "Application for <Target Role>", "body": "string"}
        if generate_email_message
        else None,
    }

    prompt_core = _build_cv_prompt(
        cv_structure,
        profile_json,
        job_description,
        prompt_override if generate_cv else None,
    )

    cover_instruction = ""
    if generate_cover_letter:
        if cover_letter_prompt and cover_letter_prompt.strip():
            cover_instruction = f"Custom cover letter instruction:\n{cover_letter_prompt.strip()}\n\n"
        else:
            cover_instruction = (
                "Cover letter requirements: professional, ATS-friendly plain text, 3-5 short paragraphs, "
                "confident concise tone, include measurable impact and a call to action.\n\n"
            )

    email_instruction = ""
    if generate_email_message:
        if email_message_prompt and email_message_prompt.strip():
            email_instruction = f"Custom email instruction:\n{email_message_prompt.strip()}\n\n"
        else:
            email_instruction = (
                "Email requirements: short professional application email, 5-8 lines max, includes greeting, "
                "1-2 key strengths, polite closing. email_message.subject should default to 'Application for <Target Role>'.\n\n"
            )

    prompt = (
        f"{prompt_core}\n\n"
        f"{cover_instruction}"
        f"{email_instruction}"
        f"Candidate Name: {full_name}\n"
        f"Target Role: {role_title}\n"
        f"Company: {company_name}\n\n"
        f"Output contract JSON shape:\n{json.dumps(output_contract, ensure_ascii=False)}\n\n"
    )

    raw = ask_gemini(prompt, model_name=model_name, api_key=gemini_api_key)
    payload = _extract_json(raw)

    sections: dict[str, Any] = {}
    if generate_cv:
        candidate_sections = payload.get("sections")
        if not isinstance(candidate_sections, dict):
            raise ValueError("AI response missing or invalid 'sections' object.")
        sections = _postprocess_sections_from_profile(candidate_sections, profile_data)

    cover_letter: str | None = None
    if generate_cover_letter:
        cover_letter = str(payload.get("cover_letter") or "").strip() or None

    email_message: Any = None
    if generate_email_message:
        email_message = payload.get("email_message")

    return {
        "sections": sections,
        "cover_letter": cover_letter,
        "email_message": email_message,
    }


def _normalize_lines(value: object) -> list[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    if isinstance(value, str) and value.strip():
        return [line.strip() for line in re.split(r"\n+", value) if line.strip()]
    return []


def _remove_years_claims(text: str) -> str:
    # Avoid overstating experience duration if the model hallucinates year counts.
    cleaned = re.sub(
        r"\b\d+(?:\.\d+)?\s*\+?\s*(?:years?|yrs?)\s*(?:of\s+)?experience\b",
        "hands-on experience",
        text,
        flags=re.IGNORECASE,
    )
    cleaned = re.sub(r"\b\d+(?:\.\d+)?\s*\+?\s*(?:years?|yrs?)\b", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned


def _strip_markdown_asterisks(text: str) -> str:
    return re.sub(r"\*+", "", text)


def _emphasize_keywords(text: str, keywords: list[str]) -> str:
    emphasized = _strip_markdown_asterisks(text)
    cleaned_keywords = sorted({k.strip() for k in keywords if isinstance(k, str) and k.strip()}, key=len, reverse=True)
    for keyword in cleaned_keywords:
        pattern = re.compile(rf"(?<!\w)({re.escape(keyword)})(?!\w)", re.IGNORECASE)
        emphasized = pattern.sub(lambda m: f"**{m.group(1)}**", emphasized)
    emphasized = re.sub(r"(~?\d+%?)", r"**\1**", emphasized)
    return emphasized


def create_pdf_from_template(
    output_path: Path | None,
    full_name: str,
    role_title: str,
    company_name: str,
    job_url: str,
    sections: dict,
) -> bytes:
    BLUE = (47, 112, 180)

    # Spacing configuration (vertical space in points for PDF)
    SECTION_SPACING = 6
    ITEM_SPACING = 2

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.add_page()

    header = sections.get("header", {}) if isinstance(sections.get("header", {}), dict) else {}

    core_skills = sections.get("core_skills", {}) if isinstance(sections.get("core_skills", {}), dict) else {}
    emphasis_keywords = []
    for group_key in [
        "languages_frameworks",
        "databases_tools",
        "testing_devops",
        "development_practices",
    ]:
        emphasis_keywords.extend(_normalize_lines(core_skills.get(group_key)))

    def draw_divider() -> None:
        y = pdf.get_y() + 1
        pdf.set_draw_color(160, 160, 160)
        pdf.set_line_width(0.4)
        pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
        pdf.ln(SECTION_SPACING)

    def write_line(text: str, h: float = 6, bold: bool = False, color: tuple[int, int, int] | None = None) -> None:
        pdf.set_x(pdf.l_margin)
        if color is not None:
            pdf.set_text_color(*color)
        else:
            pdf.set_text_color(0, 0, 0)

        if bold:
            pdf.set_font("Helvetica", "B", 11)
            pdf.multi_cell(0, h, _safe_text(text))
            pdf.set_text_color(0, 0, 0)
            return
        try:
            pdf.multi_cell(0, h, _safe_text(text), markdown=True)
        except TypeError:
            pdf.multi_cell(0, h, _safe_text(text.replace("**", "")))
        pdf.set_text_color(0, 0, 0)

    def normalize_url(raw: str) -> str:
        value = raw.strip()
        if not value:
            return value
        if value.lower().startswith(("http://", "https://", "mailto:")):
            return value
        return f"https://{value}"

    def write_clickable_line(label: str, url_text: str, link: str, h: float = 6) -> None:
        if not url_text.strip() or not link.strip():
            return
        pdf.set_x(pdf.l_margin)
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(0, 0, 0)
        label_text = _safe_text(label)
        if label_text:
            pdf.write(h, label_text)

        pdf.set_font("Helvetica", "U", 10)
        pdf.set_text_color(0, 64, 160)
        pdf.write(h, _safe_text(url_text), link=link)
        pdf.ln(h)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", "", 10)

    display_name = str(header.get("full_name") or full_name).strip()
    display_title = str(header.get("job_title") or role_title).strip()
    display_location = str(header.get("location") or "").strip()
    display_phone = str(header.get("phone") or "").strip()
    display_email = str(header.get("email") or "").strip()
    display_github = str(header.get("github") or "").strip()
    display_linkedin = str(header.get("linkedin") or "").strip()

    pdf.set_font("Helvetica", "B", 18)
    write_line(display_name, h=10, bold=True)

    if display_title:
        pdf.set_font("Helvetica", "", 12)
        write_line(display_title, h=7)

    contact_parts = [part for part in [display_location, display_phone] if part]
    if contact_parts:
        pdf.set_font("Helvetica", "", 10)
        write_line(" | ".join(contact_parts), h=6)

    if display_email:
        write_clickable_line("", display_email, f"mailto:{display_email}", h=6)

    if display_github:
        write_clickable_line("GitHub: ", display_github, normalize_url(display_github), h=6)
    if display_linkedin:
        write_clickable_line("LinkedIn: ", display_linkedin, normalize_url(display_linkedin), h=6)

    draw_divider()

    def section(title: str, lines: list[str]) -> None:
        if not lines:
            return
        pdf.set_font("Helvetica", "B", 12)
        write_line(title, h=8, bold=True, color=BLUE)
        pdf.set_font("Helvetica", "", 11)
        for line in lines:
            emphasized = _emphasize_keywords(line, emphasis_keywords)
            write_line(f"- {emphasized}", h=6)
        draw_divider()

    summary = str(sections.get("professional_summary", "")).strip()
    if summary:
        summary = _remove_years_claims(summary)
        pdf.set_font("Helvetica", "B", 12)
        write_line("Professional Summary", h=8, bold=True, color=BLUE)
        pdf.set_font("Helvetica", "", 11)
        write_line(_emphasize_keywords(summary, emphasis_keywords), h=6)
        draw_divider()

    if core_skills:
        pdf.set_font("Helvetica", "B", 12)
        write_line("Core Skills", h=8, bold=True, color=BLUE)
        pdf.set_font("Helvetica", "", 11)

        skill_groups = [
            ("Languages & Frameworks", _normalize_lines(core_skills.get("languages_frameworks"))),
            ("Databases & Tools", _normalize_lines(core_skills.get("databases_tools"))),
            ("Testing & DevOps", _normalize_lines(core_skills.get("testing_devops"))),
            ("Development Practices", _normalize_lines(core_skills.get("development_practices"))),
        ]
        for label, values in skill_groups:
            if not values:
                continue
            line = f"**{label}:** {', '.join(values)}"
            write_line(line, h=6)
        draw_divider()

    experiences = sections.get("professional_experience")
    if isinstance(experiences, list) and experiences:
        pdf.set_font("Helvetica", "B", 12)
        write_line("Professional Experience", h=8, bold=True, color=BLUE)
        for exp in experiences:
            if not isinstance(exp, dict):
                continue
            title = str(exp.get("title") or "").strip()
            company = str(exp.get("company") or "").strip()
            duration = str(exp.get("duration") or "").strip()

            heading = " ".join(part for part in [title, f"at {company}" if company else ""] if part).strip()
            if heading:
                pdf.set_font("Helvetica", "B", 11)
                write_line(heading, h=6, bold=True, color=BLUE)
            if duration:
                pdf.set_font("Helvetica", "", 10)
                write_line(duration, h=5)

            pdf.set_font("Helvetica", "", 11)
            for line in _normalize_lines(exp.get("highlights")):
                emphasized = _emphasize_keywords(line, emphasis_keywords)
                write_line(f"- {emphasized}", h=6)
            pdf.ln(ITEM_SPACING)
        draw_divider()

    projects = sections.get("personal_projects")
    if isinstance(projects, list) and projects:
        pdf.set_font("Helvetica", "B", 12)
        write_line("Personal Projects", h=8, bold=True, color=BLUE)
        for project in projects:
            if not isinstance(project, dict):
                continue
            name = str(project.get("name") or "").strip()
            tech_stack = _normalize_lines(project.get("tech_stack"))
            highlights = _normalize_lines(project.get("highlights"))

            if name:
                pdf.set_font("Helvetica", "B", 11)
                write_line(name, h=6, bold=True, color=BLUE)
            if tech_stack:
                pdf.set_font("Helvetica", "", 10)
                write_line(f"Tech Stack: {', '.join(tech_stack)}", h=5)

            pdf.set_font("Helvetica", "", 11)
            for line in highlights:
                emphasized = _emphasize_keywords(line, emphasis_keywords + tech_stack)
                write_line(f"- {emphasized}", h=6)
            pdf.ln(ITEM_SPACING)
        draw_divider()

    section("Education", _normalize_lines(sections.get("education")))
    section("Training & Certifications", _normalize_lines(sections.get("training_certifications")))

    pdf_bytes = bytes(pdf.output())

    if output_path is not None:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_bytes(pdf_bytes)

    return pdf_bytes


def _write_parsed_runs(paragraph, text: str) -> None:
    # Simple parser for **bold** markers produced by _emphasize_keywords
    parts = text.split("**")
    bold = False
    for part in parts:
        run = paragraph.add_run(part)
        run.bold = bold
        bold = not bold


def create_docx_from_template(
    output_path: Path | None,
    full_name: str,
    role_title: str,
    company_name: str,
    job_url: str,
    sections: dict,
) -> bytes:
    doc = Document()

    def _set_spacing(paragraph, before: int = 0, after: int = 6) -> None:
        try:
            paragraph.paragraph_format.space_before = Pt(before)
            paragraph.paragraph_format.space_after = Pt(after)
        except Exception:
            pass

    def add_heading(text: str, level: int = 1) -> None:
        if not text:
            return
        h = doc.add_paragraph()
        run = h.add_run(_safe_text(text))
        run.bold = True
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(12)
        else:
            run.font.size = Pt(11)
        _set_spacing(h, before=2, after=6)

    def add_paragraph_text(text: str, bold: bool = False) -> None:
        if not text:
            return
        p = doc.add_paragraph()
        if bold:
            run = p.add_run(_safe_text(text))
            run.bold = True
        else:
            _write_parsed_runs(p, _safe_text(text))
        _set_spacing(p, before=0, after=4)

    header = sections.get("header", {}) if isinstance(sections.get("header", {}), dict) else {}

    display_name = str(header.get("full_name") or full_name).strip()
    display_title = str(header.get("job_title") or role_title).strip()
    display_location = str(header.get("location") or "").strip()
    display_phone = str(header.get("phone") or "").strip()
    display_email = str(header.get("email") or "").strip()
    display_github = str(header.get("github") or "").strip()
    display_linkedin = str(header.get("linkedin") or "").strip()

    # Header
    if display_name:
        h = doc.add_paragraph()
        r = h.add_run(display_name)
        r.bold = True
        r.font.size = Pt(16)
        _set_spacing(h, before=0, after=4)
    if display_title:
        p = doc.add_paragraph(display_title)
        p.runs[0].font.size = Pt(11)
        _set_spacing(p, before=0, after=4)
    contact = " | ".join([p for p in [display_location, display_phone] if p])
    if contact:
        p = doc.add_paragraph(contact)
        _set_spacing(p, before=0, after=4)
    if display_email:
        p = doc.add_paragraph(display_email)
        _set_spacing(p, before=0, after=4)
    if display_github:
        p = doc.add_paragraph(f"GitHub: {display_github}")
        _set_spacing(p, before=0, after=4)
    if display_linkedin:
        p = doc.add_paragraph(f"LinkedIn: {display_linkedin}")
        _set_spacing(p, before=0, after=4)

    # small spacer paragraph
    p = doc.add_paragraph()
    _set_spacing(p, before=0, after=6)

    def write_section(title: str, lines: list[str]) -> None:
        if not lines:
            return
        p = doc.add_paragraph(title)
        p.runs[0].bold = True
        _set_spacing(p, before=4, after=6)
        for line in lines:
            txt = _safe_text(line)
            p = doc.add_paragraph()
            _write_parsed_runs(p, f"- {txt}")
            _set_spacing(p, before=0, after=3)

    # Professional summary
    summary = str(sections.get("professional_summary", "")).strip()
    if summary:
        p = doc.add_paragraph("Professional Summary")
        p.runs[0].bold = True
        _set_spacing(p, before=4, after=6)
        p = doc.add_paragraph()
        _write_parsed_runs(p, _emphasize_keywords(_remove_years_claims(summary), []))
        _set_spacing(p, before=0, after=4)

    core_skills = sections.get("core_skills", {}) if isinstance(sections.get("core_skills", {}), dict) else {}
    if core_skills:
        p = doc.add_paragraph("Core Skills")
        p.runs[0].bold = True
        _set_spacing(p, before=4, after=6)
        groups = [
            ("Languages & Frameworks", _normalize_lines(core_skills.get("languages_frameworks"))),
            ("Databases & Tools", _normalize_lines(core_skills.get("databases_tools"))),
            ("Testing & DevOps", _normalize_lines(core_skills.get("testing_devops"))),
            ("Development Practices", _normalize_lines(core_skills.get("development_practices"))),
        ]
        for label, vals in groups:
            if not vals:
                continue
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(", ".join(vals))
            _set_spacing(p, before=0, after=3)

    experiences = sections.get("professional_experience")
    if isinstance(experiences, list) and experiences:
        doc.add_paragraph("Professional Experience", style="Heading 2")
        for exp in experiences:
            if not isinstance(exp, dict):
                continue
            title = str(exp.get("title") or "").strip()
            company = str(exp.get("company") or "").strip()
            duration = str(exp.get("duration") or "").strip()
            heading = " ".join(part for part in [title, f"at {company}" if company else ""] if part).strip()
            if heading:
                p = doc.add_paragraph(heading)
                p.runs[0].bold = True
                _set_spacing(p, before=2, after=2)
            if duration:
                p = doc.add_paragraph(duration)
                _set_spacing(p, before=0, after=2)
            for line in _normalize_lines(exp.get("highlights")):
                p = doc.add_paragraph()
                _write_parsed_runs(p, f"- {line}")
                _set_spacing(p, before=0, after=2)

    projects = sections.get("personal_projects")
    if isinstance(projects, list) and projects:
        doc.add_paragraph("Personal Projects", style="Heading 2")
        for project in projects:
            if not isinstance(project, dict):
                continue
            name = str(project.get("name") or "").strip()
            tech_stack = _normalize_lines(project.get("tech_stack"))
            highlights = _normalize_lines(project.get("highlights"))
            if name:
                p = doc.add_paragraph(name)
                p.runs[0].bold = True
                _set_spacing(p, before=2, after=2)
            if tech_stack:
                p = doc.add_paragraph()
                p.add_run("Tech Stack: ").bold = True
                p.add_run(", ".join(tech_stack))
                _set_spacing(p, before=0, after=2)
            for line in highlights:
                p = doc.add_paragraph()
                _write_parsed_runs(p, f"- {line}")
                _set_spacing(p, before=0, after=2)

    write_section("Education", _normalize_lines(sections.get("education")))
    write_section("Training & Certifications", _normalize_lines(sections.get("training_certifications")))

    bio = BytesIO()
    doc.save(bio)
    docx_bytes = bio.getvalue()

    if output_path is not None:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_bytes(docx_bytes)

    return docx_bytes
